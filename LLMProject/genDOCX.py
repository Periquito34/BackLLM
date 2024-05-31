from docx import Document
from docx.shared import Pt

def generate_docx(requirements):
    doc = Document()
    
    for i, (requirement, priority) in enumerate(requirements, start=1):
        doc.add_heading(f'ID: HU-{i:03}', level=1)
        doc.add_paragraph(f'Nombre historia: {requirement}')
        
        doc.add_paragraph('Como usuario, quiero poder...')
        
        table = doc.add_table(rows=4, cols=2)
        
        cells = table.rows[0].cells
        cells[0].text = 'PRIORIDAD EN NEGOCIO:'
        cells[1].text = f'Prioridad: {priority}'
        
        cells = table.rows[1].cells
        cells[0].text = 'RIESGO EN DESARROLLO:'
        cells[1].text = 'Bajo'
        
        cells = table.rows[2].cells
        cells[0].text = 'PUNTOS ESTIMADOS:'
        cells[1].text = '3'
        
        cells = table.rows[3].cells
        cells[0].text = 'ITERACION ASIGNADA:'
        cells[1].text = ''
        
        doc.add_paragraph('CRITERIOS DE ACEPTACION:')
        
        criteria = [
            'El usuario puede incrementar la cantidad de unidades de un producto en el carrito.',
            'La cantidad total de unidades en el carrito se actualiza correctamente.',
            'El sistema muestra un mensaje de confirmación después de agregar las unidades.'
        ]
        
        for criterion in criteria:
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(criterion)
            run.font.size = Pt(11)
        
        doc.add_paragraph()
    
    output_path = 'generated_files/processed_requirements.docx'
    doc.save(output_path)
    
    return output_path
